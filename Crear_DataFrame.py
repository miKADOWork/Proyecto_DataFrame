# --------------------------------------------------------------------------------------------------------------------------
# Importaciones
# --------------------------------------------------------------------------------------------------------------------------

import pandas as pd
import time as tm # Paquete tiempo 
import regex
import re as re
import os
#import xlsxwriter
import openpyxl
from openpyxl.styles import PatternFill, GradientFill
import docx

# --------------------------------------------------------------------------------------------------------------------------
# Variables Globales
# --------------------------------------------------------------------------------------------------------------------------

ENCODING = "latin1"
PATH_TO_DATA = "subvenciones.csv"
NOMBRE_EXCEL_CREADO = "Resumen_Subenciones.xlsx"
NOMBRE_DIR_GUARDADO_INFORMES  = "Informes"
COLUMNA_INICIO = 5
FILA_INICIO = 5

# Para calcular el nombre de las columnas en excel
NUM_LETRAS_ALPHABET = 26                            # numero de letras en el alfabeto ingles
DECENAS = NUM_LETRAS_ALPHABET*NUM_LETRAS_ALPHABET   #
CENTENAS = DECENAS*NUM_LETRAS_ALPHABET              #
LIM = DECENAS + NUM_LETRAS_ALPHABET                 #

# --------------------------------------------------------------------------------------------------------------------------
# Colores de Celdas
# --------------------------------------------------------------------------------------------------------------------------

redFill = PatternFill   (  
                            start_color='FFFF0000',
                            end_color='FFFF0000',
                            fill_type='solid'
                        )

yellowFill = PatternFill(  
                            start_color="FFFF00",
                            fill_type='solid'
                        )

# --------------------------------------------------------------------------------------------------------------------------
# Funciones
# --------------------------------------------------------------------------------------------------------------------------


def Calculador_Nombre_Columnas (num_columna):
    """
    Args:
        num_columna (int): La possici??n de la columna de Excel numericamente (EMPEZANDO POR 0)

    Returns:
        string: La possicion con la nomenglatura de Excel para las columnas
    """
    Letra = ""
    num_veces = num_columna // NUM_LETRAS_ALPHABET - 1       # Guardamos el quocient de la division Euclidiana, le restamos uno porque trabajamos con modulo y la primera possicion seria un 1 en vez de 0
    resto = num_columna % NUM_LETRAS_ALPHABET                # Guardamos el residu de la division Euclidiana 
    
    if num_columna < NUM_LETRAS_ALPHABET:
        Letra = chr(ord("A") + num_columna)
    elif (num_columna < LIM):
        Letra = chr(ord("A") + (num_veces)) + chr(ord("A") + resto)
    elif (LIM <= num_columna):
        Letra = chr(ord("A") + (num_columna % NUM_LETRAS_ALPHABET))
        Letra = chr(ord("A") + (num_veces % NUM_LETRAS_ALPHABET)) + Letra
        num_veces = num_veces // NUM_LETRAS_ALPHABET - 1
        Letra = chr(ord("A") + num_veces) + Letra
    return Letra 


def Pinta_Columnas (rango_seleccionado, patron_relleno):
    """
    Entran un string de la forma <columna><numero_inicial>:<columna><numero_final>
    """
    
    # Separamos el rango en la celda inicial y la final
    rango = rango_seleccionado.split(sep=':')
    
    # Solo encontrara una coincidencia que se guardara en la posicion 0
    columna_nombre = re.findall("([A-Z])\w+", rango[0])[0] 

    rango = [
                rango[0].replace(columna_nombre, ""), 
                rango[1].replace(columna_nombre, "")
            ]
    
    for fila in range(int(rango[0]), int(rango[1])):
        wb.active[columna_nombre + str(fila)].fill = patron_relleno
    

def Pinta_Filas (rango_seleccionado, patron_relleno): # FALTA ACABARLA (mirar el TODO)
    """
    Entran un string de la forma <fila_inicial><numero>:<fila_final><numero>
    """
    
    columna_inicial_nombre = ""
    columna_final_nombre = ""
 
    # Separamos el rango en la celda inicial y la final 
    rango = rango_seleccionado.split(sep=':')
    columna_inicial_nombre = re.findall("([A-Z])\w+", rango[0])[0] 
    columna_final_nombre = re.findall("([A-Z])\w+", rango[1])[0] 
    
    # Calculamos la fila sobre la que trabajamos
    fila_de_trabajo = rango[0].replace(columna_inicial_nombre, "")
    fila_de_trabajo_final = re.findall("([0-9])+", rango[1])[0] 
    rango = [
                columna_inicial_nombre, 
                columna_final_nombre,
            ]
    # Falta mejorar como se realiza el rango paracontemplar casos de nom column AA ABC 
    for columna in range(ord(rango[0])-ord(rango[0]), ord(rango[1]) - ord(rango[0]) + 1):
        wb.active[Calculador_Nombre_Columnas(ord(columna_inicial_nombre) - ord("A") + columna) + fila_de_trabajo].fill = patron_relleno # Falla porque ord(columna_inicial_nombre) es mas grande que 100


def Pinta_Rango (rango_seleccionado, patron_relleno):
    """
    Args:
        rango_seleccionado (string): Continene la primera celda y la ultima del rango seleccionado
                                     El string de entrada sera de la forma siguiente:
                                        "<columna_inicial> <fila_inicial> : <columna_final> <fila final>"
                                        
        patron_relleno (format): Contiene el estilo de lo que se le aplicara a dicha seleccion de celdas
    Function:
        Funciona con una convinaci??n de PintarFilas y partes de PintarColumnas de forma que genere los strings apropiadamente para poder pasarlos a PintarFilas y
        ir realizando fila a fila los ajustes de color necesarios.
    """

    # Entra un string de la forma  "<col_inicial><fila_inicial> : <col_final><fila_final> " por tanto lo purgamos para
    # trabajar de forma mas simple
    rango = rango_seleccionado.split(sep=':')
    
    # Calculamos las possiciones de interes:
    columna_inicial = re.findall("([A-Z])\w+", rango[0])[0] 
    columna_final = re.findall("([A-Z])\w+", rango[1])[0] 
    fila_inicial = int(re.findall("[0-9]+", rango[0])[0])
    fila_final = int(re.findall("[0-9]+", rango[1])[0])
    
    # Para cada fila calculamos el string que pasaremos a la funcion Pinta_Filas
    for fila in range(fila_inicial, fila_final + 1):
        rango_para_pintar_iteracion = columna_inicial + str(fila) + ":" + columna_final + str(fila)
        Pinta_Filas(rango_para_pintar_iteracion, patron_relleno)
    

# --------------------------------------------------------------------------------------------------------------------------
# Programa
# --------------------------------------------------------------------------------------------------------------------------

# Leemos el archivo
df = pd.read_csv(   
                    PATH_TO_DATA,
                    encoding = ENCODING,
                    header = 0,
                )

# Oredenamos el dataf rame sin eliminar las repetiocones por el nombre porque es mas elegante
df = df.sort_values(by = ['Asociaci??n'])
df["Importe"] = df["Importe"].astype(float)

# Ordenamos por nombre las asociaciones para que esten alineados con los datos correctos ya que les hemos hecho un sort pero al eliminar duplicidades con un set quedan desorganizados
Lista_Asociaciones = list({a for a in df["Asociaci??n"]})
Lista_Asociaciones.sort()

dic =   {
            "Asociaci??n":   Lista_Asociaciones,
            "Max":          [float(a) for a in df.groupby(by = 'Asociaci??n',).max()["Importe"]],
            "Suma":         [float(a) for a in df.groupby(by = 'Asociaci??n',).sum()["Importe"]],
            "Media":        [round(a,2) for a in df.groupby(by = 'Asociaci??n',).mean()["Importe"]],
        }

# Machacamos la variable con los nuevos datos obtenidos y eliminamos la lista de asociaciones que ya no nos ser?? de utilidad
df = pd.DataFrame(dic)
del Lista_Asociaciones

# Escribimos estos datos en un archivo de excel 
df.to_excel (
                NOMBRE_EXCEL_CREADO,            # Creamos un libro de excel
                sheet_name = "Datos_Operados",  # Creamos una hoja destino en el libro con los datos del Dataframe
                index = False,                  # No escribe los indices de las filas
                startrow = FILA_INICIO,         # Fila de inicio del volcado del Dataframe
                startcol = COLUMNA_INICIO,      # Columna de inicio del volcado del Dataframe
            )

# Eliminamos el df porque no lo utilizamos mas pero guardamos la forma 
shape_de_df = df.shape
#del df                 # Por el momento no la eliminamos para generar los documentos

# Damos formato a la tabla que nos a generado ------------------------------------------------------------------------
# Abrimos el arxivo de excel en un Dataframe 
columnas_seleccionadas = chr(ord('A') + COLUMNA_INICIO ) + ":" + chr(ord('A') + COLUMNA_INICIO + shape_de_df[1] - 1)                    # Calculamos las letras de las columnas que seleccionamos del excel
Libro_Excel = pd.read_excel (
                                NOMBRE_EXCEL_CREADO,                                                                                    # Path al archivo en este caso el nombre porque estan en el mismo relative path
                                usecols = columnas_seleccionadas,                                                                       # Calculamos un string que contenga el rango de la tabla con la que queremos trabajar
                                header = FILA_INICIO,                                                                                   # Indicamos la primera fila con datos (que es nuestra cabezera)
                            )

# Pintamos todas las columnas
wb =  openpyxl.load_workbook(NOMBRE_EXCEL_CREADO)     # Abrimos el Libro de excel
wb.active = wb["Datos_Operados"]

# Coloreamos la cabezera de la tabla
rango_cabezera = str(chr(ord("A") + COLUMNA_INICIO)) + str(FILA_INICIO + 1) + ":" + str(chr(ord("A") + COLUMNA_INICIO + shape_de_df[1] - 1)) + str(FILA_INICIO + 1) 
Pinta_Filas(rango_cabezera, redFill)

# Coloreamos alternando con dos colores las columnas del resultado
for i in range(0, shape_de_df[1]):
    posicion = chr(ord(columnas_seleccionadas[0]) + i) + str(FILA_INICIO + 2) + ":" + chr(ord(columnas_seleccionadas[0]) + i) + str(shape_de_df[0] + FILA_INICIO + 2) # 1 porque empieza contando en 0 y excel en 1 y el otro por la linea del header 
    if (i%2 == 0):
        Pinta_Columnas(posicion, redFill)
    else:
        Pinta_Columnas(posicion, yellowFill)
        
# Ahora pintamos toda la tabla de color amarillo para testear esta funcion que selecciona rangos de la tabla
Pinta_Rango ("J6:O33", redFill)

# Ponemos en negrita en todas las celdas en las que aparezca la palaba "AMPA" dicha palabra en negrita (Usamos regex)

# Guardamos los canvios sobrescribiendo el archivo original
wb.save(NOMBRE_EXCEL_CREADO)

# Eliminamos el Excel creado
os.system(f"rm {NOMBRE_EXCEL_CREADO}")

# Por practicar mas, creamos un archivo de word que contenga estos datos en un informe a trabes del Dataframe que hemos generado previamente
if os.path.exists(f"{NOMBRE_DIR_GUARDADO_INFORMES}"):
    print("El directorio existe")
else:
    os.system(f"mkdir {NOMBRE_DIR_GUARDADO_INFORMES }")

# Nos movemos al directorio sobre el que queremos trabajar
os.chdir(f"{NOMBRE_DIR_GUARDADO_INFORMES}")

for key in df["Asociaci??n"]:
    # Creamos el documento 
    document = docx.Document()

    # Escribimos en la primera possici??n del documento
    document.add_heading    (
                                f'\t Resumen de la asociaci??n:\t {key}\n',      # String con los datos
                                3,                                              # Tamanyo (0 el mas grande a 9 el mas peque??o)
                            )

    # Creamos la tabla que contiene los datos calculados en nuestro dataframe
    table = document.add_table(rows=2, cols=3)
    
    # Guardamos la fila que sera la cabezera
    hdr_cells = table.rows[0].cells
    hdr_cells[1].text = "Resultados de la associaci??n"
    
    # Guardamos la siguiente fila en la que pondremos los titulos de cada columna
    titulos_columnas = table.rows[1].cells
    titulos_columnas[0].text = df.keys()[1]        
    titulos_columnas[1].text = df.keys()[2]  
    titulos_columnas[2].text = df.keys()[3]  
    
    # Guardamos la fila sobre la que escribiremos los datos
    row_cells = table.add_row().cells
    
    for num_col in range(0,3):
        valores = list(df.iloc[0])                              # Guardamos los datos de una fila incluiendo el nombre de la associaci??n
        row_cells[num_col].text = str(valores[num_col + 1])     # Pasamos la informaci??n a un texto
    
    # Eliminamos la primera posici??n (la 0)
    df = df.iloc[1:]                                            # machacamos la variable sin la poss. cero
    df.reset_index  ( 
                        drop=True, 
                        inplace=True,    
                    )                                           # reindexamos para que los indices empiezen por 0

    # Guardamos el documento
    document.save(f'Informe_{key}.docx')

# Volvemos al directorio de trabajo
os.chdir(f"..")

##### TODO: 
# Poner en las celdas con las asociaciones, la palabra AMPA en negrita
# Generar el informe en word con el formato deseado
# Poner diferentes separadores, que los datos no solo sean csv